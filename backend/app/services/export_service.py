from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from sqlalchemy.orm import Session

from app.models import Asset, Digest, FinalOutput, NewsCandidate, QualityCheck


def build_docx(db: Session, digest: Digest, output_path: Path) -> Path:
    doc = Document()
    doc.add_heading(f"ExTellect AI Digest — {digest.date.isoformat()}", level=1)

    image_asset = (
        db.query(Asset).filter(Asset.digest_id == digest.id, Asset.type == "image").order_by(Asset.id.desc()).first()
    )
    if image_asset and Path(image_asset.path).exists():
        doc.add_picture(image_asset.path, width=Inches(6.5))

    outputs = db.query(FinalOutput).filter(FinalOutput.digest_id == digest.id).all()
    platform_map = {o.platform: o.content for o in outputs}
    for platform in ["telegram", "max", "vk", "dzen"]:
        doc.add_heading(platform.upper(), level=2)
        doc.add_paragraph(platform_map.get(platform, ""))

    doc.add_heading("Таблица самопроверки", level=2)
    checks = db.query(QualityCheck).filter(QualityCheck.digest_id == digest.id).all()
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Проверка"
    hdr[1].text = "Статус"
    hdr[2].text = "Комментарий"
    for check in checks:
        row = table.add_row().cells
        row[0].text = check.check_name
        row[1].text = check.status
        row[2].text = check.comment

    doc.add_heading("Список источников", level=2)
    candidates = db.query(NewsCandidate).filter(NewsCandidate.digest_id == digest.id).order_by(NewsCandidate.id).all()
    for candidate in candidates:
        doc.add_paragraph(f"{candidate.title} — {candidate.source} — {candidate.url}")

    doc.add_paragraph(f"Дата генерации: {datetime.utcnow().isoformat()} UTC")
    doc.save(output_path)
    return output_path
